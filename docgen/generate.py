from docx import Document


REVIEW_DATA = {
  'quality': 'Quality of work (1-5)',
  'quantity': 'Quantity of work (1-5)',
  'initiative': 'Initiative, creativity, experience, leadership (1-5)',
  'dependability': 'Dependability and meeting commitments (1-5)',
  'interaction':
    'Interaction, supporting other team members, sharing information (1-5)',
  'meetings': 'Team meetings -- Participation, punctuality (1-5)',
  'overall': 'Overall contributions (1-5)',
  'comments':
    'Comments/feedback -- detailed, specific comments are needed here to earn '
    'full credit.'
}


def generate_peer_review_doc(reviews, destFilename, srcFilename=None):
  document = Document(srcFilename)
  paragraph = None
  format = None
  for i, review in enumerate(reviews):
    format = None
    if i == 0:
      paragraph = document.add_paragraph('Self Evaluation\n')
      format = paragraph.paragraph_format
      format.line_spacing = 1
      format.space_before = 1
      format.space_after = 1
      add_review(review, paragraph, True)
    else:
      paragraph = document.add_paragraph(f'Teammate {i} Evaluation\n')
      add_review(review, paragraph, False)

    paragraph.format = format

  if paragraph != None:
    paragraph.text += '\n'
  document.save(destFilename)


def add_review(review, paragraph, selfReview):
  for key in review.keys():
    if key == 'name':
      if selfReview:
        paragraph.text += f'\tYour Name: {review[key]}\n'
      else:
        paragraph.text += f'\tTeammate Name: {review[key]}\n'
    elif key == 'comments':
      paragraph.text += f'\t{REVIEW_DATA[key]}:\n\n\t${review[key]}\n'
    else:
      paragraph.text += f'\t{REVIEW_DATA[key]}: {review[key]}\n'
